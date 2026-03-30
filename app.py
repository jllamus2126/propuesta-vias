from flask import Flask, request, jsonify, render_template, send_file
import anthropic, os, json, sqlite3, io, zipfile, requests
from datetime import datetime, timedelta
import cloudinary, cloudinary.uploader, cloudinary.api
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pypdf import PdfReader

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 30 * 1024 * 1024

cloudinary.config(
    cloud_name=os.environ.get("CLOUDINARY_CLOUD_NAME"),
    api_key=os.environ.get("CLOUDINARY_API_KEY"),
    api_secret=os.environ.get("CLOUDINARY_API_SECRET"),
    secure=True
)

def get_db():
    db = sqlite3.connect('propuestas.db')
    db.row_factory = sqlite3.Row
    return db

def init_db():
    db = get_db()
    db.executescript('''
        CREATE TABLE IF NOT EXISTS empresas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            razon_social TEXT, nit TEXT, tipo TEXT, direccion TEXT, ciudad TEXT,
            telefono TEXT, email TEXT, rep_legal TEXT, cc_rep_legal TEXT,
            rep_es_ingeniero INTEGER DEFAULT 0, rep_matricula TEXT,
            camara_comercio TEXT, fecha_vencimiento_camara TEXT, fecha_vencimiento_rup TEXT,
            contador_nombre TEXT, contador_cc TEXT, contador_tp TEXT,
            revisor_nombre TEXT, revisor_cc TEXT, revisor_tp TEXT,
            cont_ind_nombre TEXT, cont_ind_cc TEXT, cont_ind_tp TEXT,
            capital_trabajo TEXT, patrimonio TEXT, liquidez TEXT,
            endeudamiento TEXT, rentabilidad TEXT, rentabilidad_activo TEXT,
            tiene_discapacidad INTEGER DEFAULT 0, tiene_mujeres INTEGER DEFAULT 0,
            es_mipyme INTEGER DEFAULT 0, exonerada_parafiscales INTEGER DEFAULT 0,
            fecha_constitucion TEXT, created_at TEXT DEFAULT CURRENT_TIMESTAMP
        );
        CREATE TABLE IF NOT EXISTS experiencia (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            empresa_id INTEGER, entidad TEXT, objeto TEXT, valor TEXT,
            fecha_inicio TEXT, fecha_fin TEXT, plazo TEXT, consecutivo_rup TEXT, acta TEXT,
            FOREIGN KEY (empresa_id) REFERENCES empresas(id)
        );
        CREATE TABLE IF NOT EXISTS documentos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            empresa_id INTEGER, experiencia_id INTEGER,
            tipo TEXT, nombre_archivo TEXT, url_cloudinary TEXT,
            public_id_cloudinary TEXT, fecha_subida TEXT, fecha_vencimiento TEXT,
            FOREIGN KEY (empresa_id) REFERENCES empresas(id),
            FOREIGN KEY (experiencia_id) REFERENCES experiencia(id)
        );
    ''')
    db.commit()
    db.close()

init_db()

def claude(prompt, max_tokens=2000):
    client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))
    r = client.messages.create(model="claude-sonnet-4-20250514", max_tokens=max_tokens,
        messages=[{"role": "user", "content": prompt}])
    return r.content[0].text

def pdf_texto(archivo_bytes, max_chars=15000):
    try:
        reader = PdfReader(io.BytesIO(archivo_bytes))
        t = ''
        for p in reader.pages:
            t += (p.extract_text() or '') + '\n'
            if len(t) > max_chars: break
        return t[:max_chars]
    except: return ''

def parse_json(texto):
    texto = texto.strip()
    if '```' in texto:
        for p in texto.split('```'):
            p = p.strip()
            if p.startswith('json'): p = p[4:].strip()
            if p.startswith('{'): texto = p; break
    return json.loads(texto)

def dias_hasta(f):
    try: return (datetime.strptime(f, '%Y-%m-%d') - datetime.now()).days
    except: return None

def alertas_emp(emp, docs):
    al = []
    for d in docs:
        if d.get('fecha_vencimiento'):
            dias = dias_hasta(d['fecha_vencimiento'])
            if dias is not None:
                lbl = d.get('tipo','').replace('_',' ').title()
                if dias < 0: al.append({'tipo':lbl,'msg':f'VENCIDO hace {abs(dias)} días','nivel':'rojo'})
                elif dias <= 30: al.append({'tipo':lbl,'msg':f'Vence en {dias} días','nivel':'amarillo'})
    for campo, lbl in [('fecha_vencimiento_camara','Cámara de comercio'),('fecha_vencimiento_rup','RUP')]:
        if emp.get(campo):
            dias = dias_hasta(emp[campo])
            if dias is not None:
                if dias < 0: al.append({'tipo':lbl,'msg':f'VENCIDO hace {abs(dias)} días','nivel':'rojo'})
                elif dias <= 30: al.append({'tipo':lbl,'msg':f'Vence en {dias} días','nivel':'amarillo'})
    return al

def crear_word(titulo, texto, proceso, fid):
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)
    hdr = doc.sections[0].header.paragraphs[0]
    hdr.text = f"Colombia Compra Eficiente · Res. 465/2024 V4 · {proceso} · {fid}"
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if hdr.runs: hdr.runs[0].font.size = Pt(9)
    h = doc.add_heading(titulo, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    for linea in texto.split('\n'):
        doc.add_paragraph(linea)
    ftr = doc.sections[0].footer.paragraphs[0]
    ftr.text = f"Propuestas vías · Res. 465/2024 · {datetime.now().strftime('%d/%m/%Y')}"
    ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if ftr.runs: ftr.runs[0].font.size = Pt(8)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

@app.route('/')
def index(): return render_template('index.html')

# ─── EMPRESAS ─────────────────────────────────────────────────────────────────
@app.route('/api/empresas', methods=['GET'])
def get_empresas():
    db = get_db()
    result = []
    for e in db.execute('SELECT * FROM empresas ORDER BY razon_social').fetchall():
        emp = dict(e)
        docs = [dict(d) for d in db.execute('SELECT * FROM documentos WHERE empresa_id=? AND experiencia_id IS NULL',(e['id'],)).fetchall()]
        emp['documentos'] = docs
        emp['alertas'] = alertas_emp(emp, docs)
        exp_list = []
        for x in db.execute('SELECT * FROM experiencia WHERE empresa_id=?',(e['id'],)).fetchall():
            xd = dict(x)
            xd['documentos'] = [dict(d) for d in db.execute('SELECT * FROM documentos WHERE experiencia_id=?',(x['id'],)).fetchall()]
            exp_list.append(xd)
        emp['experiencia'] = exp_list
        result.append(emp)
    db.close()
    return jsonify(result)

@app.route('/api/empresas', methods=['POST'])
def create_empresa():
    d = request.json
    db = get_db()
    cur = db.execute('''INSERT INTO empresas (razon_social,nit,tipo,direccion,ciudad,telefono,email,rep_legal,cc_rep_legal,
        rep_es_ingeniero,rep_matricula,camara_comercio,fecha_vencimiento_camara,fecha_vencimiento_rup,
        contador_nombre,contador_cc,contador_tp,revisor_nombre,revisor_cc,revisor_tp,
        cont_ind_nombre,cont_ind_cc,cont_ind_tp,capital_trabajo,patrimonio,liquidez,
        endeudamiento,rentabilidad,rentabilidad_activo,tiene_discapacidad,tiene_mujeres,
        es_mipyme,exonerada_parafiscales,fecha_constitucion) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)''',
        (d.get('razon_social',''),d.get('nit',''),d.get('tipo',''),d.get('direccion',''),d.get('ciudad',''),
         d.get('telefono',''),d.get('email',''),d.get('rep_legal',''),d.get('cc_rep_legal',''),
         d.get('rep_es_ingeniero',0),d.get('rep_matricula',''),d.get('camara_comercio',''),
         d.get('fecha_vencimiento_camara',''),d.get('fecha_vencimiento_rup',''),
         d.get('contador_nombre',''),d.get('contador_cc',''),d.get('contador_tp',''),
         d.get('revisor_nombre',''),d.get('revisor_cc',''),d.get('revisor_tp',''),
         d.get('cont_ind_nombre',''),d.get('cont_ind_cc',''),d.get('cont_ind_tp',''),
         d.get('capital_trabajo',''),d.get('patrimonio',''),d.get('liquidez',''),
         d.get('endeudamiento',''),d.get('rentabilidad',''),d.get('rentabilidad_activo',''),
         d.get('tiene_discapacidad',0),d.get('tiene_mujeres',0),d.get('es_mipyme',0),
         d.get('exonerada_parafiscales',0),d.get('fecha_constitucion','')))
    eid = cur.lastrowid; db.commit(); db.close()
    return jsonify({'id':eid,'ok':True})

@app.route('/api/empresas/<int:eid>', methods=['PUT'])
def update_empresa(eid):
    d = request.json
    db = get_db()
    db.execute('''UPDATE empresas SET razon_social=?,nit=?,tipo=?,direccion=?,ciudad=?,telefono=?,email=?,
        rep_legal=?,cc_rep_legal=?,rep_es_ingeniero=?,rep_matricula=?,camara_comercio=?,fecha_vencimiento_camara=?,
        fecha_vencimiento_rup=?,contador_nombre=?,contador_cc=?,contador_tp=?,revisor_nombre=?,revisor_cc=?,revisor_tp=?,
        cont_ind_nombre=?,cont_ind_cc=?,cont_ind_tp=?,capital_trabajo=?,patrimonio=?,liquidez=?,endeudamiento=?,
        rentabilidad=?,rentabilidad_activo=?,tiene_discapacidad=?,tiene_mujeres=?,es_mipyme=?,
        exonerada_parafiscales=?,fecha_constitucion=? WHERE id=?''',
        (d.get('razon_social',''),d.get('nit',''),d.get('tipo',''),d.get('direccion',''),d.get('ciudad',''),
         d.get('telefono',''),d.get('email',''),d.get('rep_legal',''),d.get('cc_rep_legal',''),
         d.get('rep_es_ingeniero',0),d.get('rep_matricula',''),d.get('camara_comercio',''),
         d.get('fecha_vencimiento_camara',''),d.get('fecha_vencimiento_rup',''),
         d.get('contador_nombre',''),d.get('contador_cc',''),d.get('contador_tp',''),
         d.get('revisor_nombre',''),d.get('revisor_cc',''),d.get('revisor_tp',''),
         d.get('cont_ind_nombre',''),d.get('cont_ind_cc',''),d.get('cont_ind_tp',''),
         d.get('capital_trabajo',''),d.get('patrimonio',''),d.get('liquidez',''),
         d.get('endeudamiento',''),d.get('rentabilidad',''),d.get('rentabilidad_activo',''),
         d.get('tiene_discapacidad',0),d.get('tiene_mujeres',0),d.get('es_mipyme',0),
         d.get('exonerada_parafiscales',0),d.get('fecha_constitucion',''),eid))
    db.commit(); db.close()
    return jsonify({'ok':True})

@app.route('/api/empresas/<int:eid>', methods=['DELETE'])
def delete_empresa(eid):
    db = get_db()
    for doc in db.execute('SELECT public_id_cloudinary FROM documentos WHERE empresa_id=?',(eid,)).fetchall():
        if doc['public_id_cloudinary']:
            try: cloudinary.uploader.destroy(doc['public_id_cloudinary'],resource_type='raw')
            except: pass
    db.execute('DELETE FROM documentos WHERE empresa_id=?',(eid,))
    db.execute('DELETE FROM experiencia WHERE empresa_id=?',(eid,))
    db.execute('DELETE FROM empresas WHERE id=?',(eid,))
    db.commit(); db.close()
    return jsonify({'ok':True})

@app.route('/api/empresas/<int:eid>/experiencia', methods=['POST'])
def add_experiencia(eid):
    d = request.json
    db = get_db()
    cur = db.execute('INSERT INTO experiencia (empresa_id,entidad,objeto,valor,fecha_inicio,fecha_fin,plazo,consecutivo_rup,acta) VALUES (?,?,?,?,?,?,?,?,?)',
        (eid,d.get('entidad',''),d.get('objeto',''),d.get('valor',''),d.get('fecha_inicio',''),
         d.get('fecha_fin',''),d.get('plazo',''),d.get('consecutivo_rup',''),d.get('acta','')))
    db.commit(); xid=cur.lastrowid; db.close()
    return jsonify({'id':xid,'ok':True})

@app.route('/api/experiencia/<int:xid>', methods=['DELETE'])
def delete_experiencia(xid):
    db = get_db()
    for doc in db.execute('SELECT public_id_cloudinary FROM documentos WHERE experiencia_id=?',(xid,)).fetchall():
        if doc['public_id_cloudinary']:
            try: cloudinary.uploader.destroy(doc['public_id_cloudinary'],resource_type='raw')
            except: pass
    db.execute('DELETE FROM documentos WHERE experiencia_id=?',(xid,))
    db.execute('DELETE FROM experiencia WHERE id=?',(xid,))
    db.commit(); db.close()
    return jsonify({'ok':True})

@app.route('/api/documentos/<int:doc_id>', methods=['DELETE'])
def delete_documento(doc_id):
    db = get_db()
    doc = db.execute('SELECT * FROM documentos WHERE id=?',(doc_id,)).fetchone()
    if doc and doc['public_id_cloudinary']:
        try: cloudinary.uploader.destroy(doc['public_id_cloudinary'],resource_type='raw')
        except: pass
    db.execute('DELETE FROM documentos WHERE id=?',(doc_id,))
    db.commit(); db.close()
    return jsonify({'ok':True})

# ─── SUBIR DOCUMENTO ─────────────────────────────────────────────────────────
@app.route('/api/subir-documento', methods=['POST'])
def subir_documento():
    empresa_id = request.form.get('empresa_id')
    experiencia_id = request.form.get('experiencia_id')
    tipo = request.form.get('tipo','soporte')
    extraer = request.form.get('extraer','false') == 'true'
    archivo = request.files.get('archivo')
    if not archivo: return jsonify({'error':'No se recibió archivo'}),400

    ab = archivo.read()
    try:
        res = cloudinary.uploader.upload(io.BytesIO(ab),
            folder=f"propuestas/empresa_{empresa_id or 'temp'}",
            resource_type='raw', use_filename=True, unique_filename=True)
        url=res['secure_url']; public_id=res['public_id']
    except Exception as e:
        return jsonify({'error':f'Error Cloudinary: {str(e)}'}),500

    datos={};fecha_venc=''

    if extraer and empresa_id:
        txt = pdf_texto(ab, 12000)
        def jp(schema, texto_doc, max_c=5000):
            return f"{schema}\nTexto:\n{texto_doc[:max_c]}\nSolo responde el JSON."

        prompts = {
            'camara': jp('''Extrae del texto de esta Cámara de Comercio colombiana en JSON:
{"razon_social":"","nit":"","tipo_empresa":"","direccion":"","ciudad":"","telefono":"","email":"",
"rep_legal":"","cc_rep_legal":"","revisor_nombre":"","revisor_cc":"","revisor_tp":"",
"fecha_expedicion":"YYYY-MM-DD","fecha_constitucion":"YYYY-MM-DD","es_mipyme":false}
fecha_constitucion = fecha en que se constituyó la sociedad.
fecha_expedicion = fecha en parte superior del certificado.''', txt, 10000),

            'rup': jp('''Extrae del RUP colombiano en JSON:
{"razon_social":"","nit":"","tipo_empresa":"","direccion":"","ciudad":"","telefono":"","email":"",
"rep_legal":"","cc_rep_legal":"","revisor_nombre":"","revisor_cc":"","revisor_tp":"",
"fecha_expedicion":"YYYY-MM-DD","fecha_constitucion":"YYYY-MM-DD","es_mipyme":false,
"liquidez":"","endeudamiento":"","rentabilidad":"","rentabilidad_activo":"","capital_trabajo":"","patrimonio":""}
REGLAS: liquidez=INDICE DE LIQUIDEZ, endeudamiento=INDICE DE ENDEUDAMIENTO,
rentabilidad=RENTABILIDAD DEL PATRIMONIO, rentabilidad_activo=RENTABILIDAD DEL ACTIVO,
capital_trabajo=ACTIVO CORRIENTE - PASIVO CORRIENTE, patrimonio=PATRIMONIO,
es_mipyme=true si dice PEQUEÑA/MEDIANA/MICRO.''', txt, 10000),

            'rut': jp('{"razon_social":"","nit":"","direccion":"","ciudad":"","email":"","telefono":"","fecha_constitucion":"YYYY-MM-DD"}', txt),
            'cedula': jp('{"nombre_completo":"","numero_cedula":"","fecha_expedicion":"YYYY-MM-DD"}', txt, 3000),
            'tarjeta_profesional': jp('{"nombre_completo":"","numero_tp":"","profesion":"","fecha_vencimiento":"YYYY-MM-DD"}', txt, 3000),
            'matricula_rep': jp('{"nombre_completo":"","numero_matricula":"","profesion":"","fecha_vencimiento":"YYYY-MM-DD"}', txt, 3000),
            'cert_copnia': jp('{"nombre_completo":"","numero_matricula":"","fecha_vencimiento":"YYYY-MM-DD"}', txt, 3000),
            'cedula_contador': jp('{"nombre_completo":"","numero_cedula":""}', txt, 3000),
            'tp_contador': jp('{"nombre_completo":"","numero_tp":"","fecha_vencimiento":"YYYY-MM-DD"}', txt, 3000),
            'cert_jcc_contador': jp('{"nombre_completo":"","numero_tp":"","fecha_vencimiento":"YYYY-MM-DD"}', txt, 3000),
            'cedula_revisor': jp('{"nombre_completo":"","numero_cedula":""}', txt, 3000),
            'tp_revisor': jp('{"nombre_completo":"","numero_tp":"","fecha_vencimiento":"YYYY-MM-DD"}', txt, 3000),
            'cert_jcc_revisor': jp('{"nombre_completo":"","numero_tp":"","fecha_vencimiento":"YYYY-MM-DD"}', txt, 3000),
            'cedula_cont_ind': jp('{"nombre_completo":"","numero_cedula":""}', txt, 3000),
            'tp_cont_ind': jp('{"nombre_completo":"","numero_tp":"","fecha_vencimiento":"YYYY-MM-DD"}', txt, 3000),
            'cert_jcc_cont_ind': jp('{"nombre_completo":"","numero_tp":"","fecha_vencimiento":"YYYY-MM-DD"}', txt, 3000),
            'estados_financieros': jp('''{"fecha_corte":"YYYY-MM-DD","activo_corriente":"","pasivo_corriente":"","activo_total":"",
"pasivo_total":"","patrimonio":"","utilidad_operacional":"","capital_trabajo":"","liquidez":"","endeudamiento":"","rentabilidad":"","rentabilidad_activo":""}
Calcula: capital_trabajo=activo_corriente-pasivo_corriente, liquidez=activo_corriente/pasivo_corriente,
endeudamiento=pasivo_total/activo_total, rentabilidad=utilidad_operacional/patrimonio,
rentabilidad_activo=utilidad_operacional/activo_total''', txt, 8000),
            'cert_discapacidad': jp('{"empresa":"","total_trabajadores":"","trabajadores_discapacidad":"","fecha_expedicion":"YYYY-MM-DD","fecha_vencimiento":"YYYY-MM-DD"}', txt, 4000),
            'redam': jp('{"nombre":"","cedula":"","fecha_expedicion":"YYYY-MM-DD","fecha_vencimiento":"YYYY-MM-DD","inhabilitado":false}', txt, 4000),
            'acta_accionaria': jp('{"empresa":"","porcentaje_mujeres":"","fecha_acta":"YYYY-MM-DD"}', txt, 5000),
        }

        prompt = prompts.get(tipo,'')
        if prompt:
            try:
                datos = parse_json(claude(prompt, 600))
                # Calcular vencimiento
                if tipo in ['camara','rup'] and datos.get('fecha_expedicion'):
                    try:
                        fexp = datetime.strptime(datos['fecha_expedicion'],'%Y-%m-%d')
                        fecha_venc = (fexp+timedelta(days=30)).strftime('%Y-%m-%d')
                    except: pass
                elif datos.get('fecha_vencimiento') and datos['fecha_vencimiento'] != 'YYYY-MM-DD':
                    fecha_venc = datos['fecha_vencimiento']

                # Auto-actualizar empresa
                db = get_db()
                ups, vs = [], []
                def add(col,val):
                    if val and str(val).strip() and str(val) not in ('YYYY-MM-DD','false','true',''):
                        ups.append(f'{col}=?'); vs.append(str(val))

                if tipo in ['camara','rup']:
                    add('razon_social',datos.get('razon_social')); add('nit',datos.get('nit'))
                    add('tipo',datos.get('tipo_empresa')); add('direccion',datos.get('direccion'))
                    add('ciudad',datos.get('ciudad')); add('telefono',datos.get('telefono'))
                    add('email',datos.get('email')); add('rep_legal',datos.get('rep_legal'))
                    add('cc_rep_legal',datos.get('cc_rep_legal')); add('revisor_nombre',datos.get('revisor_nombre'))
                    add('revisor_cc',datos.get('revisor_cc')); add('revisor_tp',datos.get('revisor_tp'))
                    add('fecha_constitucion',datos.get('fecha_constitucion'))
                    add('liquidez',datos.get('liquidez')); add('endeudamiento',datos.get('endeudamiento'))
                    add('rentabilidad',datos.get('rentabilidad')); add('rentabilidad_activo',datos.get('rentabilidad_activo'))
                    add('capital_trabajo',datos.get('capital_trabajo')); add('patrimonio',datos.get('patrimonio'))
                    if datos.get('es_mipyme') is True: add('es_mipyme',1)
                    if tipo=='camara' and fecha_venc: add('fecha_vencimiento_camara',fecha_venc)
                    if tipo=='rup' and fecha_venc: add('fecha_vencimiento_rup',fecha_venc)
                elif tipo=='rut':
                    add('razon_social',datos.get('razon_social')); add('nit',datos.get('nit'))
                    add('direccion',datos.get('direccion')); add('ciudad',datos.get('ciudad'))
                    add('email',datos.get('email')); add('telefono',datos.get('telefono'))
                    add('fecha_constitucion',datos.get('fecha_constitucion'))
                elif tipo=='estados_financieros':
                    add('capital_trabajo',datos.get('capital_trabajo')); add('patrimonio',datos.get('patrimonio'))
                    add('liquidez',datos.get('liquidez')); add('endeudamiento',datos.get('endeudamiento'))
                    add('rentabilidad',datos.get('rentabilidad')); add('rentabilidad_activo',datos.get('rentabilidad_activo'))
                elif tipo in ['cedula_contador','tp_contador','cert_jcc_contador']:
                    add('contador_nombre',datos.get('nombre_completo'))
                    add('contador_cc',datos.get('numero_cedula') or datos.get('numero_cc'))
                    add('contador_tp',datos.get('numero_tp'))
                elif tipo in ['cedula_revisor','tp_revisor','cert_jcc_revisor']:
                    add('revisor_nombre',datos.get('nombre_completo'))
                    add('revisor_cc',datos.get('numero_cedula') or datos.get('numero_cc'))
                    add('revisor_tp',datos.get('numero_tp'))
                elif tipo in ['cedula_cont_ind','tp_cont_ind','cert_jcc_cont_ind']:
                    add('cont_ind_nombre',datos.get('nombre_completo'))
                    add('cont_ind_cc',datos.get('numero_cedula') or datos.get('numero_cc'))
                    add('cont_ind_tp',datos.get('numero_tp'))
                elif tipo in ['matricula_rep','cert_copnia']:
                    add('rep_matricula',datos.get('numero_matricula'))

                if ups:
                    vs.append(empresa_id)
                    db.execute(f"UPDATE empresas SET {','.join(ups)} WHERE id=?",vs)
                    db.commit()
                db.close()
            except Exception as e:
                datos = {'error': str(e)}

    db = get_db()
    db.execute('INSERT INTO documentos (empresa_id,experiencia_id,tipo,nombre_archivo,url_cloudinary,public_id_cloudinary,fecha_subida,fecha_vencimiento) VALUES (?,?,?,?,?,?,?,?)',
        (empresa_id, experiencia_id or None, tipo, archivo.filename, url, public_id,
         datetime.now().strftime('%Y-%m-%d'), fecha_venc))
    db.commit(); db.close()
    return jsonify({'ok':True,'url':url,'datos_extraidos':datos,'fecha_vencimiento':fecha_venc})

# ─── ANALIZAR PLIEGO ─────────────────────────────────────────────────────────
@app.route('/api/analizar-pliego', methods=['POST'])
def analizar_pliego():
    archivo = request.files.get('pliego')
    if not archivo: return jsonify({'error':'No se recibió el pliego'}),400
    txt = pdf_texto(archivo.read(), 18000)
    prompt = f'''Analiza este pliego de condiciones colombiano (Res. 465/2024 infraestructura de transporte) y extrae en JSON:
{{"entidad":"","ciudad_entidad":"","direccion_entidad":"","numero_proceso":"","objeto":"",
"presupuesto_oficial":"","plazo_ejecucion":"","fecha_cierre":"","lote":"",
"requisitos_habilitantes":{{"financieros":{{"liquidez_minima":"","endeudamiento_maximo":"","rentabilidad_minima":"","capital_trabajo_minimo":""}},
"tecnicos":{{"experiencia_minima_valor":"","capacidad_residual_minima":""}}}},
"formatos_puntaje":{{"F7A":{{"aplica":false,"puntaje":""}},"F7B":{{"aplica":false,"puntaje":""}},"F7C":{{"aplica":false,"puntaje":""}},
"F8":{{"aplica":false,"puntaje":""}},"F9A":{{"aplica":false,"puntaje":"","tiene_bienes_relevantes":false,"bienes_relevantes":[],"solo_nacionales":true}},
"F9B":{{"aplica":false,"puntaje":""}},"F12":{{"aplica":false,"puntaje":""}},"F13":{{"aplica":false,"puntaje":""}},"F14":{{"aplica":false,"puntaje":""}}}},
"formula_consorcio":"","total_puntaje":"100"}}
Texto del pliego:
{txt}
Solo responde el JSON.'''
    try:
        resp = claude(prompt, 2000)
        return jsonify({'ok':True,'pliego':parse_json(resp)})
    except Exception as e:
        return jsonify({'error':str(e)}),500

# ─── VERIFICAR CUMPLIMIENTO ───────────────────────────────────────────────────
@app.route('/api/verificar-cumplimiento', methods=['POST'])
def verificar_cumplimiento():
    data = request.json
    pliego = data.get('pliego',{})
    emps = data.get('empresas',[])
    db = get_db()
    rf = pliego.get('requisitos_habilitantes',{}).get('financieros',{})
    res = {'empresas':[],'alertas_docs':[],'formatos_a_generar':[],'info_f9':{}}

    for ed in emps:
        e = db.execute('SELECT * FROM empresas WHERE id=?',(ed['id'],)).fetchone()
        if not e: continue
        e = dict(e)
        def chk(lbl,val,mn,mayor=True):
            try:
                v=float(str(val or '0').replace(',','.').replace('$','').replace('.','').strip())
                m=float(str(mn or '0').replace(',','.').strip())
                if m==0: return {'label':lbl,'valor':str(val or ''),'requerido':str(mn or ''),'cumple':True,'na':True}
                return {'label':lbl,'valor':str(val or ''),'requerido':str(mn or ''),'cumple':v>=m if mayor else v<=m}
            except: return {'label':lbl,'valor':str(val or ''),'requerido':str(mn or ''),'cumple':None}

        inds = [chk('Liquidez',e.get('liquidez'),rf.get('liquidez_minima')),
                chk('Endeudamiento',e.get('endeudamiento'),rf.get('endeudamiento_maximo'),False),
                chk('Rentabilidad patrimonio',e.get('rentabilidad'),rf.get('rentabilidad_minima')),
                chk('Capital de trabajo',e.get('capital_trabajo'),rf.get('capital_trabajo_minimo'))]

        docs = [dict(d) for d in db.execute('SELECT * FROM documentos WHERE empresa_id=? AND experiencia_id IS NULL',(ed['id'],)).fetchall()]
        for a in alertas_emp(e,docs):
            res['alertas_docs'].append({'empresa':e['razon_social'],**a})

        res['empresas'].append({'id':ed['id'],'razon_social':e['razon_social'],'pct':ed.get('pct',100),
            'indicadores':inds,'tiene_discapacidad':e['tiene_discapacidad'],
            'tiene_mujeres':e['tiene_mujeres'],'es_mipyme':e['es_mipyme'],'rep_es_ingeniero':e['rep_es_ingeniero']})

    fmt = pliego.get('formatos_puntaje',{})
    fmts = ['F1']
    if len(emps)>1: fmts.append('F2')
    fmts.append('F6')
    for f in ['F7A','F7B','F7C','F8']:
        if fmt.get(f,{}).get('aplica'): fmts.append(f)
    f9a=fmt.get('F9A',{}); f9b=fmt.get('F9B',{})
    if f9a.get('aplica'):
        fmts.append('F9A')
        res['info_f9']={'formato':'F9A','tiene_bienes_relevantes':f9a.get('tiene_bienes_relevantes',False),
            'bienes_relevantes':f9a.get('bienes_relevantes',[]),'solo_nacionales':f9a.get('solo_nacionales',True),
            'puntaje':f9a.get('puntaje',''),'descripcion':'Promoción de servicios nacionales o con trato nacional'}
    elif f9b.get('aplica'):
        fmts.append('F9B')
        res['info_f9']={'formato':'F9B','tiene_bienes_relevantes':False,'bienes_relevantes':[],
            'solo_nacionales':False,'puntaje':f9b.get('puntaje',''),'descripcion':'Incorporación de componente nacional en servicios extranjeros'}
    for f in ['F12','F13','F14']:
        if fmt.get(f,{}).get('aplica'): fmts.append(f)
    fmts.extend(['ANX4','F10','F11'])
    res['formatos_a_generar']=fmts
    db.close()
    return jsonify(res)

# ─── GENERAR UN FORMATO (llamado uno por uno) ─────────────────────────────────
@app.route('/api/generar-formato', methods=['POST'])
def generar_formato():
    data = request.json
    fid = data.get('formato_id')
    pl = data.get('pliego',{})
    emps = data.get('empresas_data',[])
    tipo_prop = data.get('tipo_proponente','individual')
    meta = data.get('meta',{})
    f9c = data.get('f9_config',{})
    exp_sel = data.get('experiencia_sel',[])

    e0 = emps[0] if emps else {}
    prop = e0.get('razon_social','') if tipo_prop=='individual' else f"{'CONSORCIO' if tipo_prop=='consorcio' else 'UNIÓN TEMPORAL'} {meta.get('nombre_plural','')}"
    rep = e0.get('rep_legal','') if tipo_prop=='individual' else meta.get('rep_plural','')
    cc = e0.get('cc_rep_legal','') if tipo_prop=='individual' else meta.get('cc_rep_plural','')

    base = f"""Entidad: {pl.get('entidad','')} | Ciudad: {pl.get('ciudad_entidad','')}
Proceso No.: {pl.get('numero_proceso','')} | Fecha de cierre: {pl.get('fecha_cierre','')}
Objeto: {pl.get('objeto','')}
Presupuesto: {pl.get('presupuesto_oficial','')} | Plazo: {pl.get('plazo_ejecucion','')}
Proponente: {prop} | Rep. legal: {rep} | CC: {cc}"""

    emp_str = '\n'.join([f"- {e.get('razon_social','')} NIT:{e.get('nit','')} Rep:{e.get('rep_legal','')} CC:{e.get('cc_rep_legal','')} {e.get('pct',100)}%" for e in emps])
    exp_str = '\n'.join([f"{i+1}. {x.get('entidad','')} | {x.get('objeto','')[:60]} | {x.get('valor','')} | Inicio:{x.get('fecha_inicio','')} Fin:{x.get('fecha_fin','')} | Cons.RUP:{x.get('consecutivo_rup','')}" for i,x in enumerate(exp_sel)])

    titulos = {'F1':'FORMATO 1\nCARTA DE PRESENTACIÓN DE LA PROPUESTA',
        'F2':f"FORMATO {'2A' if tipo_prop=='consorcio' else '2B'}\nCONFORMACIÓN DE PROPONENTE PLURAL",
        'F6':'FORMATO 6\nPAGOS DE SEGURIDAD SOCIAL Y APORTES LEGALES',
        'F7A':'FORMATO 7A\nPROGRAMA DE GERENCIA DE PROYECTOS',
        'F7B':'FORMATO 7B\nDISPONIBILIDAD Y CONDICIONES DE MAQUINARIA',
        'F7C':'FORMATO 7C\nPLAN DE CALIDAD',
        'F8':'FORMATO 8\nVINCULACIÓN PERSONAS CON DISCAPACIDAD',
        'F9A':'FORMATO 9A\nPUNTAJE DE INDUSTRIA NACIONAL',
        'F9B':'FORMATO 9B\nINCORPORACIÓN COMPONENTE NACIONAL',
        'F12':'FORMATO 12\nACREDITACIÓN EMPRESAS DE MUJERES',
        'F13':'FORMATO 13\nACREDITACIÓN MIPYME',
        'F14':'FORMATO 14\nCRITERIOS AMBIENTALES Y SOCIALES',
        'ANX4':'ANEXO 4\nPACTO DE TRANSPARENCIA',
        'F10':'FORMATO 10\nFACTORES DE DESEMPATE',
        'F11':'FORMATO 11\nAUTORIZACIÓN DATOS PERSONALES'}

    nombres_arch = {'F1':'Formato1_CartaPresentacion','F2':'Formato2_Consorcio_UT',
        'F6':'Formato6_SeguridadSocial','F7A':'Formato7A_GerenciaProyectos',
        'F7B':'Formato7B_Maquinaria','F7C':'Formato7C_PlanCalidad',
        'F8':'Formato8_Discapacidad','F9A':'Formato9A_IndustriaNacional',
        'F9B':'Formato9B_ComponenteNacional','F12':'Formato12_EmpresasMujeres',
        'F13':'Formato13_Mipyme','F14':'Formato14_AmbientalSocial',
        'ANX4':'Anexo4_PactoTransparencia','F10':'Formato10_Desempate',
        'F11':'Formato11_DatosPersonales'}

    def cont_str(e):
        return f"Contador:{e.get('contador_nombre','[Contador]')} CC:{e.get('contador_cc','')} TP:{e.get('contador_tp','')}\nRevisor:{e.get('revisor_nombre','No aplica')} CC:{e.get('revisor_cc','')} TP:{e.get('revisor_tp','')}\nCont.Ind:{e.get('cont_ind_nombre','No aplica')} CC:{e.get('cont_ind_cc','')} TP:{e.get('cont_ind_tp','')}\nExonerada parafiscales:{'Sí' if e.get('exonerada_parafiscales') else 'No'}"

    prompts = {
        'F1': f"Redacta FORMATO 1 - CARTA DE PRESENTACIÓN DE LA PROPUESTA completo según Res. 465/2024.\n{base}\nTipo:{'Empresa individual' if tipo_prop=='individual' else tipo_prop.upper()}\n{f'NIT:{e0.get(\"nit\",\"\")} Dir:{e0.get(\"direccion\",\"\")} Ciudad:{e0.get(\"ciudad\",\"\")} Tel:{e0.get(\"telefono\",\"\")} Email:{e0.get(\"email\",\"\")}' if tipo_prop=='individual' else emp_str}\nExperiencia:{chr(10)}{exp_str or 'Según documentos adjuntos'}\nTodos los numerales bajo juramento. Ciudad y fecha al inicio. Espacio firma al final.",
        'F2': f"Redacta FORMATO {'2A-CONSORCIO' if tipo_prop=='consorcio' else '2B-UNIÓN TEMPORAL'} completo según Res. 465/2024.\n{base}\nNombre:{meta.get('nombre_plural','')} Objeto:{meta.get('obj_plural','')} Duración:{meta.get('dur_plural','')}\nRep:{meta.get('rep_plural','')} CC:{meta.get('cc_rep_plural','')} Suplente:{meta.get('rep_suplente','')} CC:{meta.get('cc_suplente','')}\nIntegrantes:\n{emp_str}\nResponsabilidad {'solidaria' if tipo_prop=='consorcio' else 'proporcional'}. Tabla integrantes con porcentajes. Firmas de cada rep legal.",
        'F6': f"Redacta FORMATO 6 - PAGOS SEGURIDAD SOCIAL completo según Res. 465/2024 Art.50 Ley 789/2002.\n{base}\n{chr(10).join([f'Empresa:{e.get(\"razon_social\",\"\")} NIT:{e.get(\"nit\",\"\")}{chr(10)}Rep:{e.get(\"rep_legal\",\"\")} CC:{e.get(\"cc_rep_legal\",\"\")}{chr(10)}{cont_str(e)}' for e in emps])}\nDeclaración juramentada pago aportes salud, pensiones, riesgos, ICBF, SENA, cajas, FIC últimos 6 meses. Texto rep legal Y contador/revisor. Firmas al final.",
        'F7A': f"Redacta FORMATO 7A - GERENCIA DE PROYECTOS completo según Res. 465/2024.\n{base}\nDir:{e0.get('direccion','')} Email:{e0.get('email','')}\nCompromiso juramentado de implementar programa de gerencia con profesional en ingeniería o arquitectura. Firmas.",
        'F7B': f"Redacta FORMATO 7B - MAQUINARIA DE OBRA completo según Res. 465/2024.\n{base}\nCompromiso disponibilidad maquinaria requerida. Tabla equipos con condiciones funcionales. Firmas.",
        'F7C': f"Redacta FORMATO 7C - PLAN DE CALIDAD completo según Res. 465/2024.\n{base}\nCompromiso plan de calidad normas técnicas colombianas. Firmas.",
        'F8': f"Redacta FORMATO 8 - DISCAPACIDAD completo según Res. 465/2024.\n{base}\nCertificación total trabajadores y personas con discapacidad. Referencia certificado Ministerio de Trabajo. Tabla y firmas.",
        'F9A': f"Redacta FORMATO 9A - INDUSTRIA NACIONAL completo según Res. 465/2024.\n{base}\nConfig:{json.dumps(f9c,ensure_ascii=False)}\n{'Bienes relevantes:'+','.join(f9c.get('bienes_relevantes',[])) if f9c.get('tiene_bienes_relevantes') else 'Sin bienes relevantes específicos en este proceso.'}\nOfrecimiento apoyo industria nacional. Tabla bienes. Firmas.",
        'F9B': f"Redacta FORMATO 9B - COMPONENTE NACIONAL EXTRANJEROS completo según Res. 465/2024.\n{base}\nProponente extranjero incorporando componente nacional. Tabla y firmas.",
        'F12': f"Redacta FORMATO 12 - EMPRESAS DE MUJERES completo según Res. 465/2024.\n{base}\nAcreditación empresa de mujeres normativa vigente. Declaración y firmas.",
        'F13': f"Redacta FORMATO 13 - MIPYME completo según Res. 465/2024 Ley 590/2000.\n{base}\nAcreditación micro/pequeña/mediana empresa. Declaración y firmas.",
        'F14': f"Redacta FORMATO 14 - CRITERIOS AMBIENTALES Y SOCIALES completo según Res. 465/2024 Decreto 142/2023.\n{base}\nCompromiso programa ambiental y social numeral 4.2.4 documento base. Firmas.",
        'ANX4': f"Redacta ANEXO 4 - PACTO DE TRANSPARENCIA completo según Res. 465/2024.\n{base}\nTodos los compromisos i) al xv): cumplir ley, buena fe, no falsificación, libre competencia, no colusión, no sobornos, preguntas escritas, lealtad, compostura audiencias. Firma al final.",
        'F10': f"Redacta FORMATO 10 - FACTORES DE DESEMPATE completo según Res. 465/2024.\n{base}\nMipyme:{'Sí' if any(e.get('es_mipyme') for e in emps) else 'No'} Mujeres:{'Sí' if any(e.get('tiene_mujeres') for e in emps) else 'No'} Discapacidad:{'Sí' if any(e.get('tiene_discapacidad') for e in emps) else 'No'}\nDeclaración factores desempate aplicables. Firmas.",
        'F11': f"Redacta FORMATO 11 - DATOS PERSONALES completo según Res. 465/2024 Ley 1581/2012.\n{base}\nTitulares:{', '.join([e.get('rep_legal','') or e.get('razon_social','') for e in emps])}\nAutorización tratamiento datos personales. Firmas.",
    }

    prompt = prompts.get(fid,'')
    if not prompt: return jsonify({'error':f'Formato {fid} no reconocido'}),400

    try: texto = claude(prompt, 2500)
    except Exception as e: return jsonify({'error':str(e)}),500

    proceso = pl.get('numero_proceso','proceso').replace('/','_')
    word = crear_word(titulos.get(fid,fid), texto, proceso, fid)
    return send_file(io.BytesIO(word),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True, download_name=f"{nombres_arch.get(fid,fid)}_{proceso}.docx")

# ─── DESCARGAR ZIP DOCUMENTOS ─────────────────────────────────────────────────
@app.route('/api/descargar-docs-zip', methods=['POST'])
def descargar_docs_zip():
    data = request.json
    emps = data.get('empresas',[])
    exp_ids = data.get('experiencia_ids',[])
    proceso = data.get('numero_proceso','proceso').replace('/','_')
    docs_ing = data.get('docs_ingeniero',[])
    db = get_db()
    carp = {'camara':'Habilitantes','rup':'Habilitantes','cedula':'Habilitantes','rut':'Habilitantes',
        'redam':'Habilitantes','estados_financieros':'Habilitantes','capacidad_residual':'Habilitantes',
        'matricula_rep':'Habilitantes','cert_copnia':'Habilitantes','cedula_contador':'Habilitantes',
        'tp_contador':'Habilitantes','cert_jcc_contador':'Habilitantes','cedula_revisor':'Habilitantes',
        'tp_revisor':'Habilitantes','cert_jcc_revisor':'Habilitantes','cedula_cont_ind':'Habilitantes',
        'tp_cont_ind':'Habilitantes','cert_jcc_cont_ind':'Habilitantes',
        'cert_discapacidad':'Puntaje','acta_accionaria':'Puntaje'}
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf,'w',zipfile.ZIP_DEFLATED) as zf:
        for ed in emps:
            e = db.execute('SELECT * FROM empresas WHERE id=?',(ed['id'],)).fetchone()
            if not e: continue
            ne = e['razon_social'].replace(' ','_')[:20]
            for doc in db.execute('SELECT * FROM documentos WHERE empresa_id=? AND experiencia_id IS NULL',(ed['id'],)).fetchall():
                if doc['url_cloudinary']:
                    try:
                        r = requests.get(doc['url_cloudinary'],timeout=20)
                        if r.ok:
                            c = carp.get(doc['tipo'],'Otros')
                            zf.writestr(f"{proceso}/{c}/{ne}_{doc['tipo']}_{doc['nombre_archivo']}",r.content)
                    except: pass
            for xid in exp_ids:
                x = db.execute('SELECT * FROM experiencia WHERE id=? AND empresa_id=?',(xid,ed['id'])).fetchone()
                if not x: continue
                for doc in db.execute('SELECT * FROM documentos WHERE experiencia_id=?',(xid,)).fetchall():
                    if doc['url_cloudinary']:
                        try:
                            r = requests.get(doc['url_cloudinary'],timeout=20)
                            if r.ok:
                                zf.writestr(f"{proceso}/Habilitantes/Experiencia/{ne}_{x['entidad'][:15].replace(' ','_')}_{doc['nombre_archivo']}",r.content)
                        except: pass
        for d in docs_ing:
            if d.get('url'):
                try:
                    r = requests.get(d['url'],timeout=20)
                    if r.ok: zf.writestr(f"{proceso}/Habilitantes/IngenieroAvalador_{d.get('tipo','')}_{d.get('nombre','doc')}",r.content)
                except: pass
    db.close()
    zip_buf.seek(0)
    return send_file(zip_buf,mimetype='application/zip',as_attachment=True,
        download_name=f"Documentos_{proceso}.zip")

if __name__ == '__main__':
    port = int(os.environ.get('PORT',5000))
    app.run(host='0.0.0.0',port=port,debug=False)

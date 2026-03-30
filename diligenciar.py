"""
Motor de diligenciamiento de formatos oficiales CCE Res. 465/2024 V4
Reemplaza campos [entre corchetes] con datos reales del proponente y el proceso
"""
from docx import Document
from docx.oxml.ns import qn
import copy, re, io
from datetime import datetime

def reemplazar_en_parrafo(parrafo, reemplazos):
    """Reemplaza texto en un párrafo preservando formato"""
    texto_completo = ''.join([r.text for r in parrafo.runs])
    texto_nuevo = texto_completo
    for viejo, nuevo in reemplazos.items():
        texto_nuevo = texto_nuevo.replace(viejo, str(nuevo))
    if texto_nuevo != texto_completo and parrafo.runs:
        parrafo.runs[0].text = texto_nuevo
        for r in parrafo.runs[1:]:
            r.text = ''

def reemplazar_en_doc(doc, reemplazos):
    """Reemplaza en todos los párrafos y tablas del documento"""
    for p in doc.paragraphs:
        reemplazar_en_parrafo(p, reemplazos)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    reemplazar_en_parrafo(p, reemplazos)
    # También en encabezados y pies de página
    for section in doc.sections:
        for p in section.header.paragraphs:
            reemplazar_en_parrafo(p, reemplazos)
        for p in section.footer.paragraphs:
            reemplazar_en_parrafo(p, reemplazos)

def campos_base(pliego, proponente, rep_legal, cc_rep, tipo_prop):
    """Campos comunes a todos los formatos"""
    hoy = datetime.now().strftime('%d de %B de %Y').replace(
        'January','enero').replace('February','febrero').replace('March','marzo').replace(
        'April','abril').replace('May','mayo').replace('June','junio').replace(
        'July','julio').replace('August','agosto').replace('September','septiembre').replace(
        'October','octubre').replace('November','noviembre').replace('December','diciembre')
    return {
        '[Número del Proceso de Contratación]': pliego.get('numero_proceso',''),
        '[NOMBRE DE LA ENTIDAD]': pliego.get('entidad','').upper(),
        '[Dirección de la entidad]': pliego.get('direccion_entidad',''),
        '[Dirección de la Entidad]': pliego.get('direccion_entidad',''),
        '[Ciudad]': pliego.get('ciudad_entidad',''),
        '[Incluir número del proceso de contratación]': pliego.get('numero_proceso',''),
        '[Incluir número del Proceso de Contratación]': pliego.get('numero_proceso',''),
        '[Indicar proceso de contratación]': pliego.get('numero_proceso',''),
        '[Incluir cuando el proceso es estructurado por lotes o grupos]': pliego.get('lote','') or '',
        '[Incluir cuando el proceso sea estructurado por lotes o grupos]': pliego.get('lote','') or '',
        '[Incluir cuando el Proceso de Contratación es estructurado por lotes o grupos]': pliego.get('lote','') or '',
        '[Incluir cuando el Proceso de Contratación sea estructurado por lotes o grupos]': pliego.get('lote','') or '',
        '[Indicar el lote o lotes a los cuales se presenta oferta]': pliego.get('lote','') or '',
        '[Indicar el lote o lotes a los cuales se presenta oferta.]': pliego.get('lote','') or '',
    }

def diligenciar_F1(template_path, pliego, empresa, tipo_prop, meta=None, exp_sel=None):
    doc = Document(template_path)
    e = empresa[0] if isinstance(empresa, list) else empresa
    rep = e.get('rep_legal','')
    cc = e.get('cc_rep_legal','')
    nit = e.get('nit','')
    razon = e.get('razon_social','')
    dir_emp = e.get('direccion','')
    ciudad_emp = e.get('ciudad','')
    tel = e.get('telefono','')
    email = e.get('email','')
    matricula = e.get('rep_matricula','') or ''
    hoy = datetime.now().strftime('%d de %B de %Y')

    if tipo_prop == 'individual':
        nombre_prop = razon
        rep_str = rep
    else:
        nombre_prop = f"{'CONSORCIO' if tipo_prop=='consorcio' else 'UNIÓN TEMPORAL'} {meta.get('nombre_plural','')}"
        rep_str = meta.get('rep_plural', rep)

    reemplazos = campos_base(pliego, nombre_prop, rep_str, cc, tipo_prop)
    reemplazos.update({
        '[Nombre del representante legal del proponente]': rep_str,
        '[Nombre del proponente]': nombre_prop,
        '[Nombre del proponente- persona natural]': nombre_prop,
        '[Incluir para procesos de contratación adelantados por SECOP II]': '',
        '[Marque con una X la característica aplica al proponente]': 'X',
        '[Incluir los folios o el número de folios donde se encuentra la información reservada]': '',
        '[Incluir la norma que le otorga el carácter de reservado]': '',
        '[En caso de que el proceso de contratación se adelante a través del SECOP II deberá incluirse lo siguiente:]': '',
    })
    reemplazar_en_doc(doc, reemplazos)
    # Llenar campos de firma al final
    for p in doc.paragraphs:
        if 'Nombre del Proponente' in p.text and '___' in p.text:
            p.runs[0].text = f"Nombre del Proponente: {nombre_prop}"
        if 'Nombre del representante legal' in p.text and '___' in p.text:
            p.runs[0].text = f"Nombre del representante legal: {rep_str}"
        if 'C. C. No.' in p.text and '___' in p.text:
            p.runs[0].text = f"C. C. No. {cc}"
        if 'Dirección' in p.text and '___' in p.text and 'correo' not in p.text.lower():
            p.runs[0].text = f"Dirección: {dir_emp}"
        if 'Correo electrónico' in p.text and '___' in p.text:
            p.runs[0].text = f"Correo electrónico: {email}"
        if 'Ciudad' in p.text and p.text.strip().startswith('Ciudad'):
            p.runs[0].text = f"Ciudad: {ciudad_emp}"
        if 'Matr' in p.text and 'cula' in p.text and matricula:
            p.runs[0].text = f"Matrícula Profesional: {matricula}"
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

def diligenciar_F2(template_path, pliego, empresas, tipo_prop, meta):
    doc = Document(template_path)
    reemplazos = campos_base(pliego, meta.get('nombre_plural',''), meta.get('rep_plural',''), meta.get('cc_rep_plural',''), tipo_prop)

    # Reemplazar integrantes
    for i, e in enumerate(empresas[:2], 1):
        reemplazos[f'[nombre del representante legal del integrante {i}]'] = e.get('rep_legal','')
        reemplazos[f'[nombre o razón social del integrante {i}]'] = e.get('razon_social','')

    reemplazos['[indicar el nombre]'] = meta.get('rep_plural','')
    reemplazos['[Definir los eventos en los cuales puede intervenir el representante suplente del consorcio.]'] = 'cuando el representante principal no pueda actuar'
    reemplazos['[a]'] = ''
    reemplazos['[El Proponente deberá diligenciar el Formato dependiendo de la forma asociativa con la que se presente (consorcio, unión temporal)]'] = ''

    reemplazar_en_doc(doc, reemplazos)

    # Llenar tabla de integrantes y porcentajes
    for table in doc.tables:
        for row in table.rows:
            texto_fila = ' '.join([c.text for c in row.cells])
            if 'integrantes del consorcio' in texto_fila.lower() or 'Indicar los nombres' in texto_fila:
                for i, e in enumerate(empresas):
                    if i < len(row.cells):
                        row.cells[i].paragraphs[0].runs[0].text = e.get('razon_social','') if row.cells[i].paragraphs[0].runs else ''

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

def diligenciar_F6(template_path, pliego, empresa, tipo_prop):
    doc = Document(template_path)
    e = empresa[0] if isinstance(empresa, list) else empresa
    razon = e.get('razon_social','')
    nit = e.get('nit','')
    rep = e.get('rep_legal','')
    cc_rep = e.get('cc_rep_legal','')
    rev_nombre = e.get('revisor_nombre','') or e.get('contador_nombre','')
    rev_cc = e.get('revisor_cc','') or e.get('contador_cc','')
    rev_tp = e.get('revisor_tp','') or e.get('contador_tp','')
    exonerada = bool(e.get('exonerada_parafiscales'))
    ciudad_camara = e.get('camara_comercio','Bogotá D.C')
    hoy = datetime.now()
    fecha_firma = f"{hoy.day} días del mes de {hoy.strftime('%B').replace('March','marzo').replace('February','febrero').replace('January','enero').replace('April','abril').replace('May','mayo').replace('June','junio').replace('July','julio').replace('August','agosto').replace('September','septiembre').replace('October','octubre').replace('November','noviembre').replace('December','diciembre')} de {hoy.year}"

    reemplazos = campos_base(pliego, razon, rep, cc_rep, tipo_prop)
    reemplazos.update({
        '[Incluir el nombre del representante legal de la persona jurídica]': rep,
        '[Incluir el nombre del representante legal de la persona jurídica o el revisor fiscal, según corresponda]': rev_nombre or rep,
        '[Incluir el número de identificación]': rev_cc or cc_rep,
        '[Incluir la razón social de la persona jurídica]': razon,
        '[Incluir el NIT]': nit,
        '[Incluir el nombre del revisor fiscal, según corresponda]': rev_nombre,
        '[Incluir número de tarjeta profesional]': rev_tp,
        '[fecha de constitución]': e.get('fecha_constitucion',''),
        '[En el evento en que la sociedad no tenga más de seis (6) meses de constituida, deberá acreditar los pagos a partir de la fecha de su constitución como se indica a continuación:]': '',
        '[Cuando la persona jurídica no esté exonerada en el pago al sistema de aportes parafiscales, deberá incluir el siguiente texto y ajustar el formato en lo\t correspondiente:]': '',
        '[Cuando la persona jurídica esté exonerada de aportes parafiscales de acuerdo con el artículo 114-1 del Estatuto Tributario modificado por el artículo 65 de la Ley 1918 de 2016, deberá incluir el siguiente texto y ajustar el formato en lo correspondiente]': '',
        '[Cuando la persona jurídica no haya tenido personal a cargo dentro de los seis (6) meses anteriores a la presentación de la propuesta deberá manifestarlo de la siguiente manera:]': '',
        '[En caso de presentar acuerdo de pago con alguna de las entidades anteriormente mencionadas, se deberá precisar el valor y el plazo previsto para el acuerdo de pago, con indicación del cumplimiento de esta obligación, caso en el cual deberá anexar copia del acuerdo de pago correspondiente y el comprobante de pago soporte del mes anterior al cierre del procedimiento de contratación]': '',
        '[Este formato debe ser diligenciado por las personas jurídicas nacionales y las extranjeras con domicilio o sucursal en Colombia las cuales deberán acreditar este requisito respecto del personal vinculado en Colombia]': '',
    })
    reemplazar_en_doc(doc, reemplazos)

    # Llenar campos de firma
    for p in doc.paragraphs:
        txt = p.text
        if rev_nombre and ('________' in txt or '______' in txt) and not p.runs:
            continue
        if 'Ciudad' in txt and '_____' in txt:
            if p.runs: p.runs[0].text = f"En constancia, se firma en {e.get('ciudad','Bogotá D.C')} a los {fecha_firma}."

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

def diligenciar_simple(template_path, pliego, empresa, tipo_prop, meta=None):
    """Para F7, F8, F9, F10, F11, F12, F13, F14, ANX4 — reemplazo básico"""
    doc = Document(template_path)
    e = empresa[0] if isinstance(empresa, list) else empresa
    razon = e.get('razon_social','')
    rep = e.get('rep_legal','')
    cc = e.get('cc_rep_legal','')
    nit = e.get('nit','')
    rev = e.get('revisor_nombre','') or e.get('contador_nombre','')
    rev_cc = e.get('revisor_cc','') or e.get('contador_cc','')

    if meta and tipo_prop != 'individual':
        nombre_prop = f"{'CONSORCIO' if tipo_prop=='consorcio' else 'UNIÓN TEMPORAL'} {meta.get('nombre_plural','')}"
        rep_usar = meta.get('rep_plural', rep)
        cc_usar = meta.get('cc_rep_plural', cc)
    else:
        nombre_prop = razon
        rep_usar = rep
        cc_usar = cc

    reemplazos = campos_base(pliego, nombre_prop, rep_usar, cc_usar, tipo_prop)
    reemplazos.update({
        '[Nombre del Proponente - persona jurídica]': nombre_prop,
        '[Nombre del Proponente- persona natural]': nombre_prop,
        '[Nombre del Proponente - persona natural]': nombre_prop,
        '[ Nombre del Proponente-persona natural]': nombre_prop,
        '[ Nombre del Proponente - persona natural]': nombre_prop,
        '[Nombre del representante del Proponente Plural]': rep_usar,
        '[Incluir el nombre de la persona natural, el representante legal de la persona jurídica o el revisor fiscal, según corresponda]': rev or rep_usar,
        '[Incluir el nombre del representante legal de la persona jurídica o del revisor fiscal, según corresponda]': rev or rep_usar,
        '[Incluir el nombre del representante legal de la persona jurídica y del revisor fiscal, cuando este último exista de acuerdo con los requerimientos de ley, o el contador]': f"{rep_usar} y {rev}" if rev else rep_usar,
        '[Incluir el nombre del representante legal y el contador o revisor fiscal, si están obligados a tenerlo]': f"{rep_usar} y {rev}" if rev else rep_usar,
        '[Incluir el nombre de la persona natural y el contador público]': rep_usar,
        '[Incluir el número de identificación]': cc_usar,
        '[Incluir los números de identificación]': f"C.C. {cc_usar}" + (f" y C.C. {rev_cc}" if rev_cc else ''),
        '[Incluir la razón social de la persona jurídica]': razon,
        '[Indicar nombre del Proponente]': nombre_prop,
        '[Incluir el NIT]': nit,
        '[Indicar si actúa como representante legal o revisor fiscal]': 'Representante Legal',
        '[Indicar si actúa como, persona natural (nombre propio), representante legal o revisor fiscal]': 'Representante Legal',
        '[Indicar si actúa como representante legal o revisor fiscal]': 'Representante Legal',
        '[incluir el nombre del establecimiento de comercio]': razon,
        '[Indicar si es micro, pequeña y mediana empresa]': 'Pequeña empresa',
        # Instrucciones entre corchetes — eliminar
        '[Este formato se diligencia por el representante legal o el revisor fiscal, según corresponda, de la persona jurídica en el que mayoritariamente participen mujeres cabeza de familia y/o mujeres víctima de violencia intrafamiliar. En el evento que la oferta la presente una entidad privada sin ánimo de lucro, ya sea fundación, corporación o asociación, se ajustará el formato en lo pertinente. La información aquí vertida contiene datos sensibles, los cuales están sujetos a reserva legal y, por tanto, no podrán publicarse en el SECOP I y II para su conocimiento]': '',
        '[Este Formato debe diligenciarse por los Proponentes personas jurídicas o los integrantes que sean personas jurídicas del Proponente Plural cuyo porcentaje de participación sea al menos del diez por ciento (10 %) y acrediten la condición de emprendimientos y empresas de mujeres.]': '',
        '[Este Formato ÚNICAMENTE debe ser diligenciado por los Proponentes nacionales o extranjeros con trato nacional, o los Proponentes Plurales integrados por estos. En ningún caso el Formato debe diligenciarse por los Proponentes extranjeros sin derecho a trato nacional que opten por el puntaje correspondiente a la incorporación de componente nacional en servicios extranjeros]': '',
        '[El presente Formato lo diligenciará toda persona (Proponente, socia o trabajador) que al momento de presentar su oferta contenga datos sensibles, para que la Entidad Estatal contratante garantice el tratamiento adecuado a estos datos]': '',
        '[El proponente para acreditar el número de personas con discapacidad en su planta de personal, deberá aportar el certificado expedido por el Ministerio de Trabajo, el cual deberá estar vigente a la fecha de cierre del Proceso de Contratación.]': '',
        '[El Proponente escogerá una de las siguientes dos (2) opciones para acreditar la condición de emprendimiento y empresa de mujer:]': '',
        '[Opción 1. Incorporar si la participación accionaria de la persona jurídica en su mayoría son mujeres y los derechos de propiedad han pertenecido a éstas durante el último año.]': '',
        '[Opción 2. Incorporar si por lo menos el cincuenta por ciento (50 %) de los empleos del nivel directivo de la persona jurídica son ejercidos por mujeres y estas han estado vinculadas laboralmente durante al menos el último año.]': '',
        '[Indicar la fecha (día/mes/año) desde que se cumple con dicha condición]': '',
        '[Opción 1. Incorporar si la Entidad Estatal determina que existe al menos un bien nacional relevante contenido en el Registro de Productores de Bienes Nacionales, aplicando la definición de Servicios Nacionales del artículo 2.2.1.1.1.3.1 del Decreto 1082 de 2015 y la metodología definida en la Matriz 4- Bienes nacionales relevantes para la obra pública del sector transporte]': '',
        '[Opción 2. Incorporar si la Entidad determina que no existen bienes nacionales relevantes incluidos en el Registro de Productores de Bienes Nacionales, aplicando la definición de Servicios Nacionales del artículo 2.2.1.1.1.3.1 del Decreto 1082 de 2015]': '',
        '[Opción 3. Esta opción ÚNICAMENTE puede ser diligenciada por los proponentes extranjeros con derecho a trato nacional o Proponentes Plurales conformados por estos, que manifiesten su voluntad de acogerse a la regla de origen de su país]': '',
        '[el Proponente incluirá el porcentaje definido por la Entidad Estatal en el numeral 4.3.1 del documento base que sea por lo menos del cuarenta por ciento (40 %), sin perjuicio de incluir uno superior]': '40%',
        '[Incluir el nombre de la Entidad]': pliego.get('entidad',''),
        '[Incluir el nombre de la Entidad Estatal]': pliego.get('entidad',''),
        '[Incluir nombre de la Entidad]': pliego.get('entidad',''),
        '[incluir número de atención de la Entidad]': '',
        '[incluir el correo electrónico]': '',
        '[incluir página web de la Entidad]': '',
        '[lunes a viernes de 8:00 am a 6:00 p.m.]': 'lunes a viernes de 8:00 am a 6:00 p.m.',
        '[Incluir el link en donde se encuentra la Política de Tratamiento de Datos Personales]': '',
        '[Incluir el nombre del representante legal de la persona jurídica]': rep_usar,
        '[Nombre y firma del representante legal de la persona jurídica o el revisor fiscal, según corresponda]': rep_usar,
        '[Nombre y firma del representante legal]': rep_usar,
        '[Nombre y firma del contador público]': rev or '',
        '[Nombre y firma del contador o revisor fiscal si está obligado a tenerlo]': rev or '',
        '[Nombre y firma de la persona natural]': rep_usar,
        '[Nombre y firma de la persona natural, el representante legal de la persona jurídica o el revisor fiscal, según corresponda]': rev or rep_usar,
        '[Tratándose de Proponentes Plurales, este formato lo presentará el integrante o los integrantes que tengan una participación de por lo menos el veinticinco por ciento (25 %) en el Consorcio o en la Unión Temporal y aporten mínimo el veinticinco por ciento (25 %) de la experiencia acreditada en la oferta]': '',
        '[Firma del Proponente o de su representante legal]': '',
        '[Firma del Proponente o de su representante Legal]': '',
        '[El interesado persona natural que solicita limitar la convocatoria del Proceso de Contratación a Mipyme, acreditará la condición de Mipyme con el diligenciamiento de este Formato y entregará los documentos requeridos en el artículo 2.2.1.2.4.2.4. del Decreto 1082 de 2015. En todo caso, cuando la calidad de Mipyme se acredite con la presentación del RUP vigente y en firme, no será necesario diligenciar este Formato ni aportar los mencionados documentos]': '',
        '[El interesado persona jurídica que solicita limitar la convocatoria del Proceso de Contratación a Mipyme, acreditará la condición de Mipyme con el diligenciamiento de este Formato y entregará los documentos requeridos en el artículo 2.2.1.2.4.2.4. del Decreto 1082 de 2015. En todo caso, cuando la calidad de Mipyme se acredite con la presentación del RUP vigente y en firme, no será necesario diligenciar este Formato ni aportar los mencionados documentos]': '',
        '[y revisor fiscal, en caso de estar obligado según los requerimientos de ley, o el contador]': f"y {rev}" if rev else '',
    })

    # Campos de firma con datos reales
    for p in doc.paragraphs:
        t = p.text
        if 'Nombre del Proponente' in t and ('_____' in t or ':' in t):
            if p.runs: p.runs[0].text = f"Nombre del Proponente: {nombre_prop}"
        elif 'Nombre del representante legal' in t and ('_____' in t or ':' in t):
            if p.runs: p.runs[0].text = f"Nombre del representante legal: {rep_usar}"
        elif 'C. C. No.' in t and '_____' in t:
            if p.runs: p.runs[0].text = f"C. C. No. {cc_usar}"
        elif 'Dirección de correo' in t and '_____' in t:
            if p.runs: p.runs[0].text = f"Dirección de correo: {e.get('email','')}"
        elif 'Ciudad' in t and t.strip().startswith('Ciudad') and '_____' in t:
            if p.runs: p.runs[0].text = f"Ciudad: {e.get('ciudad','')}"
        elif 'Matrícula' in t and '_____' in t and e.get('rep_matricula'):
            if p.runs: p.runs[0].text = f"Matrícula Profesional: {e.get('rep_matricula','')}"

    reemplazar_en_doc(doc, reemplazos)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

def diligenciar_formato(fid, template_path, pliego, empresas, tipo_prop, meta=None, exp_sel=None, f9_config=None):
    """Punto de entrada principal"""
    emp = empresas[0] if empresas else {}
    try:
        if fid == 'F1':
            return diligenciar_F1(template_path, pliego, emp, tipo_prop, meta, exp_sel)
        elif fid == 'F2':
            return diligenciar_F2(template_path, pliego, empresas, tipo_prop, meta)
        elif fid == 'F6':
            return diligenciar_F6(template_path, pliego, emp, tipo_prop)
        else:
            return diligenciar_simple(template_path, pliego, emp, tipo_prop, meta)
    except Exception as ex:
        # Si falla, retornar el template sin modificar
        with open(template_path, 'rb') as f:
            return f.read()

print("Módulo diligenciar.py creado OK")
